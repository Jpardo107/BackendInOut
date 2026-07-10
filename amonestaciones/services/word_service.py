from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_cell_border_top(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:color"), "333333")
    borders.append(top)


def generar_word_amonestacion(amonestacion):
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    lines = amonestacion.carta.replace("\r\n", "\n").split("\n")
    signature_start = next(
        (index for index, line in enumerate(lines) if "EMPLEADOR" in line.upper() and index >= len(lines) // 2),
        None,
    )
    body_lines = lines[:signature_start] if signature_start is not None else lines
    while body_lines and (not body_lines[-1].strip() or set(body_lines[-1].strip()) <= {"_", "-", " "}):
        body_lines.pop()

    for raw_line in body_lines:
        line = raw_line.strip()
        paragraph = document.add_paragraph()
        if not line:
            paragraph.paragraph_format.space_after = Pt(2)
            continue
        run = paragraph.add_run(line)
        if line.upper() in {"SEÑOR", "PRESENTE"} or line.startswith("Ref.:"):
            run.bold = True
        if line.startswith(("•", "- ")):
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_paragraph().paragraph_format.space_after = Pt(20)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell in table.rows[0].cells:
        cell.width = Cm(7.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_border_top(cell)

    employer, worker = table.rows[0].cells
    employer.text = "EMPLEADOR\nINOUT SEGURIDAD SpA\nRUT 76.435.221-1"
    worker.text = (
        f"TRABAJADOR\n{amonestacion.persona.nombre_completo.upper()}\n"
        f"RUT {amonestacion.persona.rut}"
    )
    for cell in (employer, worker):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(10)
                if paragraph == cell.paragraphs[0]:
                    run.bold = True

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output

from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from documentacion.services.r2_storage import download_document_to_fileobj
from reportes.models import ReporteInforme


BRAND = "193040"
LIGHT_BORDER = "DEE4E8"
CONTENT_WIDTH_CM = 16.6


def _shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or "Sin información"))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_width(table, width_cm=CONTENT_WIDTH_CM):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def _set_table_borders(table, color=LIGHT_BORDER, size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        borders.append(element)


def _set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def _heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(BRAND)
    return paragraph


def _body(document, text):
    paragraph = document.add_paragraph(str(text or "Sin información."))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(7)
    return paragraph


def _info_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    _set_table_width(table)
    _set_table_borders(table)
    header = table.rows[0].cells
    _shade(header[0], BRAND)
    _shade(header[1], BRAND)
    _set_cell_text(header[0], "Campo", bold=True, color="FFFFFF", size=10)
    _set_cell_text(header[1], "Detalle", bold=True, color="FFFFFF", size=10)
    for cell, width in zip(header, (4.2, 12.4)):
        cell.width = Cm(width)
        _set_cell_margins(cell)
    for label, value in rows:
        cells = table.add_row().cells
        for cell, width in zip(cells, (4.2, 12.4)):
            cell.width = Cm(width)
            _set_cell_margins(cell)
        _shade(cells[0], "F1F4F6")
        _set_cell_text(cells[0], label, bold=True)
        _set_cell_text(cells[1], value)
    return table


def _police_absence_banner(document):
    table = document.add_table(rows=1, cols=1)
    _set_table_width(table)
    cell = table.cell(0, 0)
    cell.width = Cm(CONTENT_WIDTH_CM)
    _set_cell_margins(cell, top=180, bottom=180)
    _shade(cell, BRAND)
    _set_cell_text(
        cell,
        "Personal policial no presente en el procedimiento",
        bold=True,
        color="FFFFFF",
        size=10,
    )
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def _items(value):
    if not value:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _item_text(item):
    if isinstance(item, dict):
        return " | ".join(f"{key.replace('_', ' ').title()}: {value}" for key, value in item.items())
    return str(item)


def _list_section(document, title, value):
    items = _items(value)
    if not items:
        return
    _heading(document, title, level=2)
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(_item_text(item))


def _add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def _add_image_block(document, image, index, include_recommendation):
    container = document.add_table(rows=1, cols=1)
    _set_table_width(container)
    _set_table_borders(container, color="D5DADF", size="7")
    container_cell = container.cell(0, 0)
    container_cell.width = Cm(CONTENT_WIDTH_CM)
    _set_cell_margins(container_cell, top=160, start=180, bottom=160, end=180)
    _shade(container_cell, "FBFCFD")
    title_paragraph = container_cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = Pt(5)
    title = title_paragraph.add_run(f"Fotografía {index}")
    title.bold = True
    title.font.name = "Aptos Display"
    title.font.size = Pt(11)
    title.font.color.rgb = RGBColor.from_string(BRAND)

    table = container_cell.add_table(rows=1, cols=2)
    _set_table_width(table, width_cm=15.7)
    image_cell, text_cell = table.rows[0].cells
    image_cell.width = Cm(7.7)
    text_cell.width = Cm(8)
    _set_cell_margins(image_cell, top=60, start=60, bottom=60, end=120)
    _set_cell_margins(text_cell, top=80, start=120, bottom=80, end=60)
    row_pr = table.rows[0]._tr.get_or_add_trPr()
    row_pr.append(OxmlElement("w:cantSplit"))

    try:
        buffer = BytesIO()
        download_document_to_fileobj(image.storage_key, buffer)
        buffer.seek(0)
        image_cell.paragraphs[0].add_run().add_picture(buffer, width=Cm(7.5), height=Cm(5.5))
    except Exception:
        _set_cell_text(image_cell, "Imagen no disponible")

    _set_cell_text(text_cell, image.descripcion or "Sin descripción.")
    if include_recommendation:
        paragraph = text_cell.add_paragraph()
        title = paragraph.add_run("Recomendación\n")
        title.bold = True
        title.font.color.rgb = RGBColor.from_string(BRAND)
        paragraph.add_run(image.recomendacion_usuario or "Sin recomendación.")

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def generar_word_reporte(reporte):
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(45, 45, 45)
    normal.paragraph_format.line_spacing = 1.1

    top_line = document.add_paragraph()
    top_line.paragraph_format.space_after = Pt(4)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:color"), BRAND)
    border.append(bottom)
    top_line._p.get_or_add_pPr().append(border)

    title = "PRE-INFORME" if reporte.tipo_reporte == ReporteInforme.TIPO_PRE_INFORME else "INFORME DE VULNERABILIDADES"
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor.from_string(BRAND)
    subtitle = document.add_paragraph("Generado por INOUT Seguridad SpA")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(9)

    _info_table(document, [
        ("Tipo", "Pre-informe" if reporte.tipo_reporte == ReporteInforme.TIPO_PRE_INFORME else "Informe de vulnerabilidades"),
        ("Zona", reporte.zona or "Sin información"),
        ("Instalación", reporte.instalacion.nombre),
        ("Fecha de emisión", reporte.fecha_emision.strftime("%d-%m-%Y")),
        ("Autor", reporte.autor_nombre or str(reporte.usuario_creador)),
        ("Cargo", reporte.autor_cargo or "No informado"),
    ])

    _heading(document, "Descripción de los hechos")
    _body(document, reporte.descripcion_hechos)

    if reporte.tipo_reporte == ReporteInforme.TIPO_PRE_INFORME:
        _heading(document, "Personal policial")
        if reporte.personal_policial_presente:
            _info_table(document, [
                ("Presente", "Sí"),
                ("Personal presente", reporte.personal_presente or "No informado"),
                ("Carabinero a cargo", reporte.carabinero_cargo or "No informado"),
                ("Patente de patrulla", reporte.patente_patrulla or "No informada"),
                ("Número de carro policial", reporte.numero_carro_policial or "No informado"),
            ])
        else:
            _police_absence_banner(document)
    else:
        _heading(document, "Análisis IA del informe")
        if reporte.criticidad_general:
            _info_table(document, [("Criticidad general", reporte.criticidad_general.title())])
        if reporte.resumen_ejecutivo:
            _heading(document, "Resumen ejecutivo", level=2)
            _body(document, reporte.resumen_ejecutivo)
        _list_section(document, "Riesgos detectados", reporte.riesgos_detectados)
        _list_section(document, "Matriz de riesgo", reporte.matriz_riesgo)
        _list_section(document, "Recomendaciones IA", reporte.recomendaciones_ia)
        if reporte.conclusion_profesional:
            _heading(document, "Conclusión profesional", level=2)
            _body(document, reporte.conclusion_profesional)
        if reporte.texto_final_pdf:
            _heading(document, "Texto final", level=2)
            _body(document, reporte.texto_final_pdf)

    images = reporte.imagenes.all().order_by("orden", "id")
    if images:
        document.add_page_break()
        _heading(document, "Registro fotográfico")
        for index, image in enumerate(images, 1):
            _add_image_block(document, image, index, reporte.tipo_reporte == ReporteInforme.TIPO_VULNERABILIDADES)

    if reporte.tipo_reporte == ReporteInforme.TIPO_VULNERABILIDADES and reporte.analisis_final_usuario:
        _heading(document, "Análisis final")
        _body(document, reporte.analisis_final_usuario)

    _add_page_number(section.footer.paragraphs[0])
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output

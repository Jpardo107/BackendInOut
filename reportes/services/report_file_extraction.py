import os

from rest_framework.exceptions import ValidationError


SUPPORTED_REPORT_EXTENSIONS = {".pdf", ".docx", ".txt"}
SUPPORTED_REPORT_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


def validate_report_file(file, max_size):
    _, ext = os.path.splitext(getattr(file, "name", "") or "")
    ext = ext.lower()
    content_type = getattr(file, "content_type", "")

    if ext not in SUPPORTED_REPORT_EXTENSIONS:
        raise ValidationError({"archivo": "Formato no permitido. Usa PDF, DOCX o TXT."})

    if content_type and content_type not in SUPPORTED_REPORT_CONTENT_TYPES:
        raise ValidationError({"archivo": "Tipo de archivo no permitido. Usa PDF, DOCX o TXT."})

    if getattr(file, "size", 0) > max_size:
        raise ValidationError({"archivo": "El informe supera el tamano maximo permitido."})


def extract_text_from_report_file(file):
    _, ext = os.path.splitext(getattr(file, "name", "") or "")
    ext = ext.lower()
    file.seek(0)

    if ext == ".pdf":
        return _extract_pdf_text(file)
    if ext == ".docx":
        return _extract_docx_text(file)
    if ext == ".txt":
        return _extract_txt_text(file)

    raise ValidationError({"archivo": "Formato no soportado para extraccion de texto."})


def _extract_pdf_text(file):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValidationError({"archivo": "El servidor no tiene soporte para leer PDF."}) from exc

    try:
        reader = PdfReader(file)
        parts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts).strip()
    except Exception as exc:
        raise ValidationError({"archivo": "No se pudo extraer texto del PDF."}) from exc


def _extract_docx_text(file):
    try:
        from docx import Document
    except ImportError as exc:
        raise ValidationError({"archivo": "El servidor no tiene soporte para leer DOCX."}) from exc

    try:
        document = Document(file)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    table_cells.append(" | ".join(values))
        return "\n\n".join([*paragraphs, *table_cells]).strip()
    except Exception as exc:
        raise ValidationError({"archivo": "No se pudo extraer texto del DOCX."}) from exc


def _extract_txt_text(file):
    raw = file.read()
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValidationError({"archivo": "No se pudo leer el archivo TXT."})
